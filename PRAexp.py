# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'form.ui'
# Author : Tejas Achar
# Sorry the code is messy, But Works
# Created by: PyQt5 UI code generator 5.14.2
#
# WARNING! All changes made in this file will be lost!

import time
import nltk
import heapq
from PyQt5.QtCore import QThread, pyqtSignal
from PyQt5.QtWidgets import (QApplication, QDialog,
                             QProgressBar, QPushButton)
from PyQt5 import QtCore, QtGui, QtWidgets
from docx import Document
import urllib.request
import time
from googlesearch import search
from bs4 import BeautifulSoup
import re
from docx.shared import Inches
import re
import requests
import docx
import os
from PIL import Image
from docx.opc.constants import RELATIONSHIP_TYPE as RT
nltk.download('punkt')
nltk.download('stopwords')



class Ui_test(object):

    NumOfImages = [0, 1, 2]
    def GenerateDetailed(self):
      if self.lineEdit_2.text() == "":
        self.lineEdit_2.setPlaceholderText("This field cannot be empty")
      else:
        self.progressbar.show()
        self.progressbar.setValue(7)
        AdditionalUrls = []
        keyword = self.lineEdit_2.text()
        print(keyword)
        List_of_file_Names = []
        for j in search(keyword,tld="com",num = 10, stop=1, pause=2):
            url = j
            res = requests.get(url)
            html_page = res.content
            soup = BeautifulSoup(html_page, 'html.parser')
            text = soup.find_all('p')

            output = ''
            blacklist = [
                    '[document]',
                    'noscript',
                    'header',
                    'html',
                    'meta',
                    'head',
                    'input',
                    'script',
                    'style'
                    # there may be more elements you don't want, such as "style", etc.
            ]

        for t in text:
              if t.parent.name not in blacklist:
                     output += '{} '.format(t.getText())

        for j in search(keyword, tld="com", num=10, stop=5, pause=2):
                AdditionalUrls.append(j)


        def download_baidu(keyword):
            self.progressbar.setValue(11)
            url = 'https://image.baidu.com/search/flip?tn=baiduimage&ie=utf-8&word=' + keyword + " images" + '&ct=201326592&v=flip'
            result = requests.get(url)
            html = result.text
            pic_url = re.findall('"objURL":"(.*?)",', html, re.S)
            i = 0

            for each in pic_url[:3]:
                x = 33
                print(each)
                self.progressbar.setValue(x)
                try:
                    pic = requests.get(each, timeout=10)
                    self.progressbar.setValue(x+33)
                except requests.exceptions.ConnectionError:
                    print('exception')
                    continue
                fileType = each[each.rindex('.') + 1:]
                if "?" in fileType:
                    ActualFile = fileType[:fileType.find("?")]
                    string = keyword + '_' + str(i) + "." + ActualFile
                else:
                    string = keyword + '_' + str(i) + "." + fileType
                if fileType != "jpeg":
                    List_of_file_Names.append(string)
                fp = open(string, 'wb')
                fp.write(pic.content)
                fp.close()
                i += 1
                x += 33
        download_baidu(keyword)

        document = Document()
        document.add_heading(keyword.upper(),0)

        p = document.add_paragraph()
        r = p.add_run()
        r2 = p.add_run()
        r3 = p.add_run()
        r.add_break()


        for j in List_of_file_Names:
            print(j)
            r.add_picture(j, width=Inches(2.0), height=Inches(1.5))
        r.add_break()
        r.add_text(output)
        r.add_break()
        r.add_break()
        r.add_break()
        r.add_break()
        r2.bold = True
        r2.add_text("Additional Links :")

        for i in AdditionalUrls:
            r3.add_break()
            r3.add_break()
            r3.add_text(i)
            r3.add_break()

        document.save(keyword+'.docx')
        self.progressbar.setValue(100)
        file = keyword+".docx"
        os.startfile(file)




    def GenerateShort(self):
          if self.lineEdit_2.text() == "":
            self.lineEdit_2.setPlaceholderText("This field cannot be empty")
          else:
            self.progressbar.show()
            self.progressbar.setValue(7)
            AdditionalUrls = []
            keyword = self.lineEdit_2.text()
            print(keyword)
            List_of_file_Names = []
            for j in search(keyword,tld="com",num = 10, stop=1, pause=2):
                url = j
                res = requests.get(url)
                html_page = res.content
                soup = BeautifulSoup(html_page, 'html.parser')
                text = soup.find_all('p')

                output = ''
                blacklist = [
                        '[document]',
                        'noscript',
                        'header',
                        'html',
                        'meta',
                        'head',
                        'input',
                        'script',
                        'style'
                        # there may be more elements you don't want, such as "style", etc.
                ]

            for t in text:
                  if t.parent.name not in blacklist:
                         output += '{} '.format(t.getText())

            for j in search(keyword, tld="com", num=10, stop=5, pause=2):
                    AdditionalUrls.append(j)


            article_text = re.sub(r'\[[0-9]*\]', ' ', output)
            article_text = re.sub(r'\s+', ' ', output)

            # Removing special characters and digits
            formatted_article_text = re.sub('[^a-zA-Z]', ' ', article_text )
            formatted_article_text = re.sub(r'\s+', ' ', formatted_article_text)

            sentence_list = nltk.sent_tokenize(article_text)

            stopwords = nltk.corpus.stopwords.words('english')

            word_frequencies = {}
            for word in nltk.word_tokenize(formatted_article_text):
                        if word not in stopwords:
                            if word not in word_frequencies.keys():
                                word_frequencies[word] = 1
                            else:
                                word_frequencies[word] += 1

            maximum_frequncy = max(word_frequencies.values())

            for word in word_frequencies.keys():
                        word_frequencies[word] = (word_frequencies[word]/maximum_frequncy)

            sentence_scores = {}
            for sent in sentence_list:
                        for word in nltk.word_tokenize(sent.lower()):
                            if word in word_frequencies.keys():
                                if len(sent.split(' ')) < 30:
                                    if sent not in sentence_scores.keys():
                                        sentence_scores[sent] = word_frequencies[word]
                                    else:
                                        sentence_scores[sent] += word_frequencies[word]
            summary_sentences = heapq.nlargest(7, sentence_scores, key=sentence_scores.get)

            summary = ' '.join(summary_sentences)
            print("++++++++++++++++++++++++++++++++++++++++++",summary)


            def download_baidu(keyword):
                self.progressbar.setValue(11)
                url = 'https://image.baidu.com/search/flip?tn=baiduimage&ie=utf-8&word=' + keyword + " images" + '&ct=201326592&v=flip'
                result = requests.get(url)
                html = result.text
                pic_url = re.findall('"objURL":"(.*?)",', html, re.S)
                i = 0

                for each in pic_url[:3]:
                    x = 33
                    print(each)
                    self.progressbar.setValue(x)
                    try:
                        pic = requests.get(each, timeout=10)
                        self.progressbar.setValue(x+33)
                    except requests.exceptions.ConnectionError:
                        print('exception')
                        continue
                    fileType = each[each.rindex('.') + 1:]
                    if "?" in fileType:
                        ActualFile = fileType[:fileType.find("?")]
                        string = keyword + '_' + str(i) + "." + ActualFile
                    else:
                        string = keyword + '_' + str(i) + "." + fileType
                    if fileType != "jpeg":
                        List_of_file_Names.append(string)
                    fp = open(string, 'wb')
                    fp.write(pic.content)
                    fp.close()
                    i += 1
                    x += 33
            download_baidu(keyword)

            document = Document()
            document.add_heading(keyword.upper(),0)

            p = document.add_paragraph()
            r = p.add_run()
            r2 = p.add_run()
            r3 = p.add_run()
            r.add_break()


            for j in List_of_file_Names:
                print(j)
                r.add_picture(j, width=Inches(2.0), height=Inches(1.5))
            r.add_break()
            r.add_text(summary)
            r.add_break()
            r.add_break()
            r.add_break()
            r.add_break()
            r2.bold = True
            r2.add_text("Additional Links :")

            for i in AdditionalUrls:
                r3.add_break()
                r3.add_break()
                r3.add_text(i)
                r3.add_break()

            document.save(keyword+'.docx')
            self.progressbar.setValue(100)
            file = keyword+".docx"
            os.startfile(file)

    def setupUi(self, test):
        test.setObjectName("test")
        test.setFixedSize(691, 375)

        self.gridLayoutWidget = QtWidgets.QWidget(test)
        # self.progressbar = QtWidgets.QProgressBar(test)
        self.gridLayoutWidget.setGeometry(QtCore.QRect(90, 80, 521, 210))
        self.gridLayoutWidget.setObjectName("gridLayoutWidget")
        self.gridLayout = QtWidgets.QGridLayout(self.gridLayoutWidget)
        self.gridLayout.setContentsMargins(0, 0, 0, 0)
        self.gridLayout.setObjectName("gridLayout")
        self.Detailed_Gen = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.Detailed_Gen.setObjectName("Detailed_Gen")
        self.gridLayout.addWidget(self.Detailed_Gen, 1, 2, 1, 1)
        self.lineEdit_2 = QtWidgets.QLineEdit(self.gridLayoutWidget)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.gridLayout.addWidget(self.lineEdit_2, 0, 0, 1, 3)
        self.Short_Gen = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.Short_Gen.setObjectName("Short_Gen")
        self.gridLayout.addWidget(self.Short_Gen, 1, 1, 1, 1)
        self.Clear_Button = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.Clear_Button.setObjectName("Clear_Button")
        self.gridLayout.addWidget(self.Clear_Button, 1, 0, 1, 1)
        # self.gridLayout.addWidget(self.progressbar,2,2,2,2)
        self.verticalLayoutWidget = QtWidgets.QWidget(test)
        self.verticalLayoutWidget.setGeometry(QtCore.QRect(90, 240, 521, 31))
        self.verticalLayoutWidget.setObjectName("verticalLayoutWidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.verticalLayoutWidget)
        self.verticalLayout.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout.setObjectName("verticalLayout")
        self.progressbar = QtWidgets.QProgressBar(self.verticalLayoutWidget)
        self.verticalLayout.addWidget(self.progressbar)
        self.progressbar.hide()
        # self.progress = QtWidgets.QProgressBar(self.progressbar)

        self.Clear_Button.clicked.connect(self.lineEdit_2.clear)
        QtCore.QMetaObject.connectSlotsByName(test)

        #TEsting
        self.verticalLayoutWidget_2 = QtWidgets.QWidget(test)
        self.verticalLayoutWidget_2.setGeometry(QtCore.QRect(270, 280, 171, 41))
        self.verticalLayoutWidget_2.setObjectName("verticalLayoutWidget_2")


        self.label = QtWidgets.QLabel(test)
        self.label.setGeometry(QtCore.QRect(300, 350, 121, 16))
        self.label.setObjectName("label")
        self.label.setText("kleverme.com | 2020")
        self.label_2 = QtWidgets.QLabel(test)
        self.label_2.setGeometry(QtCore.QRect(110, 30, 471, 41))
        font = QtGui.QFont()
        font.setFamily("Yu Gothic UI Semibold")
        font.setPointSize(22)
        font.setBold(True)
        font.setWeight(75)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.label_2.setText("Personal Research Assistant")
        self.label_3 = QtWidgets.QLabel(test)
        self.label_3.setGeometry(QtCore.QRect(260, 80, 191, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.label_3.setText("Experimental Phase")
        self.retranslateUi(test)






    def retranslateUi(self, test):
        _translate = QtCore.QCoreApplication.translate
        test.setWindowTitle(_translate("test", "test"))
        self.Detailed_Gen.setText(_translate("test", "Generate Short"))
        self.lineEdit_2.setPlaceholderText(_translate("test", "Type here"))
        self.Short_Gen.setText(_translate("test", "Generate Detailed"))
        self.Clear_Button.setText(_translate("test", "Clear"))
        self.Short_Gen.clicked.connect(self.GenerateDetailed)
        self.Detailed_Gen.clicked.connect(self.GenerateShort)

        self.label.setText(_translate("test", "kleverme.com | 2020"))
        self.label_2.setText(_translate("test", "Personal Research Assistant"))
        self.label_3.setText(_translate("test", "Experimental Stage"))



if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    test = QtWidgets.QWidget()
    ui = Ui_test()
    ui.setupUi(test)
    test.show()
    sys.exit(app.exec_())
